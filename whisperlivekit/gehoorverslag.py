"""Gehoorverslag-export -- docx-generatie vanaf een sessietranscript.

Bouwt een Word-document uit een reeds samengevoegde, chronologisch
gesorteerde transcript-lijst (zelfde vorm als de `channel_id=all`-tak van
`GET /sessions/{id}/transcript` in `TriviasServer.py`): een letterlijke,
chronologische weergave van vragen en antwoorden, hoormedewerker-tekst
cursief t.o.v. overige rollen, nooit content weglaten.

Dit is bewust een EXPORT, geen rapport-generator: geen classificatie naar
secties, geen samenvatting, geen structuur die het transcript niet al had.
De hoormedewerker plakt de uitvoer zelf in het door INDiGO gegenereerde
"Rapport nader gehoor" en voegt daar zelf structuur/tekstblokken aan toe --
zie features/gehoorverslag-automatisering.md in de projectrepo voor de
volledige onderbouwing van deze scope-keuze.

Puur, van FastAPI losgekoppeld: elke functie is los te unit-testen met
synthetische segment-lijsten (zelfde patroon als cross_channel_gate.py).
Wijzigt nooit de input-lijst.
"""

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

logger = logging.getLogger("whisperlivekit.gehoorverslag")

# Python-poort van web_trivias/app.js's ROLES/getRoleLabel() -- houd deze
# twee plekken bewust in sync als de frontend-rollen ooit wijzigen.
ROLE_LABELS: Dict[str, str] = {
    "employee": "Hoormedewerker",
    "interpreter": "Tolk",
    "lawyer": "Gemachtigde",
    "foreign": "Vreemdeling",
    "default": "Spreker",
}

DEFAULT_PAUSE_GAP_THRESHOLD_MS = 10_000
FONT_NAME = "Verdana"
FONT_SIZE_PT = 9


def channel_id_to_role(channel_id: str) -> str:
    """Python-poort van app.js's channelIdToRoleId(): alle foreign_<taal>-
    kanalen vallen samen tot rol "foreign", overige kanalen blijven zichzelf."""
    if not channel_id:
        return "default"
    if channel_id.startswith("foreign"):
        return "foreign"
    return channel_id


def role_label(channel_id: str) -> str:
    role = channel_id_to_role(channel_id)
    return ROLE_LABELS.get(role, role)


def _format_timestamp(ms: Optional[int]) -> str:
    if ms is None:
        return "??:??"
    total_seconds = int(ms) // 1000
    h, rem = divmod(total_seconds, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _set_base_style(document: "Document") -> None:
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)


def _add_styled_paragraph(document: "Document", text: str, *, bold: bool = False, italic: bool = False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.bold = bold
    run.italic = italic
    return p


def _add_voorblad(
    document: "Document",
    session_id: str,
    ordered_segments: List[Dict[str, Any]],
    session_meta: Optional[Dict[str, Any]],
) -> None:
    _add_styled_paragraph(document, "Rapport van gehoor", bold=True)
    _add_styled_paragraph(document, "Automatisch gegenereerd door Trivias STT -- zie opmerking onderaan dit voorblad.")

    channels_seen: Dict[str, str] = {}
    for seg in ordered_segments:
        ch = seg.get("channel_id")
        if ch and ch not in channels_seen:
            channels_seen[ch] = role_label(ch)

    first_ms = ordered_segments[0].get("start_ms") if ordered_segments else None
    last_ms = ordered_segments[-1].get("end_ms") if ordered_segments else None
    session_meta = session_meta or {}
    languages = session_meta.get("languages", {})

    table = document.add_table(rows=0, cols=2)

    def _row(label: str, value: Optional[str]) -> None:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value or "(niet vastgelegd -- handmatig aanvullen)"

    _row("Sessie-ID", session_id)
    _row("Datum", session_meta.get("date"))
    _row("Aanvangstijd (relatief)", _format_timestamp(first_ms))
    _row("Eindtijd (relatief)", _format_timestamp(last_ms))
    for ch, label in channels_seen.items():
        _row(f"Kanaal -- {label}", languages.get(ch, ch))
    # Velden die WI 2021/13 als vast onderdeel van het voorblad noemt, maar
    # die Trivias nergens vastlegt (client-side, of helemaal niet ingevoerd)
    # -- expliciet als invulplaceholder, nooit verzonnen.
    _row("Naam hoormedewerker (geslacht)", None)
    _row("Naam tolk (geslacht)", None)
    _row("Tolk-registratieniveau (of reden niet-registertolk)", None)
    _row("Naam/aanwezigheid gemachtigde of hulpverlener (VWN)", None)
    _row("Termijn voor correcties en aanvullingen", None)

    _add_styled_paragraph(
        document,
        "Let op: dit document is automatisch gegenereerd vanaf de opgenomen "
        "audio en is een letterlijke weergave, geen beoordeling van "
        "geloofwaardigheid en geen beslissing. Sectie-indeling hieronder (indien "
        "aanwezig) is een classificatie ter ondersteuning, geen oordeel -- "
        "controleer en corrigeer waar nodig. Velden hierboven die niet "
        "automatisch zijn ingevuld ('handmatig aanvullen') dienen te worden "
        "aangevuld voordat dit rapport als officieel rapport van gehoor "
        "wordt gebruikt. Pauze-annotaties in dit document zijn een "
        "benadering (afgeleid uit tijdsgaten tussen segmenten), geen harde "
        "stilte-detectie.",
    )
    document.add_page_break()


def _render_segment(document: "Document", seg: Dict[str, Any]) -> None:
    channel_id = seg.get("channel_id", "") or ""
    label = role_label(channel_id)
    text = (seg.get("text_final") or "").strip()
    is_hoormedewerker = channel_id_to_role(channel_id) == "employee"
    _add_styled_paragraph(
        document,
        f"[{_format_timestamp(seg.get('start_ms'))}] {label}: {text}",
        italic=is_hoormedewerker,
    )


def _render_flat_body(
    document: "Document",
    ordered_segments: List[Dict[str, Any]],
    pause_gap_threshold_ms: int,
) -> None:
    """Platte, puur chronologische weergave -- de fallback zonder
    classificatie (geen LLM-backend geconfigureerd, of classificatie leverde
    niets bruikbaars op). Dit was v1's enige rendering en blijft ongewijzigd
    beschikbaar."""
    _add_styled_paragraph(document, "Verloop van het gehoor", bold=True)

    prev_end_ms: Optional[int] = None
    for seg in ordered_segments:
        start_ms = seg.get("start_ms")
        if prev_end_ms is not None and start_ms is not None:
            gap = start_ms - prev_end_ms
            if gap >= pause_gap_threshold_ms:
                _add_styled_paragraph(
                    document,
                    f"Opmerking rapporteur: [pauze van ongeveer {gap // 1000}s "
                    "-- afgeleid uit tijdsverloop tussen segmenten, geen harde detectie]",
                    italic=True,
                )
        _render_segment(document, seg)
        end_ms = seg.get("end_ms")
        prev_end_ms = end_ms if end_ms is not None else start_ms


def order_segments_for_report(merged_segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Filter (geen lege tekst) + chronologisch sorteren."""
    return sorted(
        (s for s in merged_segments if (s.get("text_final") or "").strip()),
        key=lambda s: s.get("start_ms") if s.get("start_ms") is not None else 0,
    )


def build_gehoorverslag_docx(
    session_id: str,
    merged_segments: List[Dict[str, Any]],
    session_meta: Optional[Dict[str, Any]] = None,
    pause_gap_threshold_ms: int = DEFAULT_PAUSE_GAP_THRESHOLD_MS,
) -> "Document":
    """Bouwt het Word-document: voorblad + platte chronologische body.

    merged_segments: lijst van dicts met minimaal channel_id, start_ms,
    end_ms, text_final (zelfde vorm als de merged-transcript-helper in
    TriviasServer.py). Segmenten zonder (niet-lege) text_final worden
    overgeslagen -- er is dan geen content om te rapporteren -- maar elk
    segment MET tekst wordt altijd opgenomen, nooit gefilterd op
    vertrouwen/rol: audio is authoritative (CLAUDE.md), niet iets waar deze
    functie een oordeel over mag hebben.
    """
    ordered = order_segments_for_report(merged_segments)

    document = Document()
    _set_base_style(document)

    _add_voorblad(document, session_id, ordered, session_meta)
    _render_flat_body(document, ordered, pause_gap_threshold_ms)

    return document
